#!/usr/bin/env python3
"""Insert generated PNGs into the thesis DOCX at Figure caption paragraphs.

Usage:
  python insert_figures_into_docx.py
  python insert_figures_into_docx.py --docx path/to/thesis.docx --figures out
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent

# Map figure id → PNG filename
FIG_MAP = {
    "2.1": "Fig_2_1_PRISMA.png",
    "3.1": "Fig_3_1_framework_pipeline.png",
    "3.2": "Fig_3_2_layer_map.png",
    "3.3": "Fig_3_3_HybridBlock.png",
    "3.4": "Fig_3_4_DAPABlock.png",
    "4.0": "Fig_4_0_dataset_class_counts.png",
    "4.1": "Fig_4_1_training_curves_mAP50.png",
    "4.2": "Fig_4_2_perclass_mAP50.png",
    "4.4": "Fig_4_4_GFLOPs_by_imgsz.png",
}


def set_run_font(run, size=11, italic=True):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.italic = italic


def insert_picture_before(paragraph, image_path: Path, width_in=5.8):
    """Insert a centered paragraph with the picture immediately before `paragraph`."""
    new_p = OxmlElement("w:p")
    paragraph._element.addprevious(new_p)
    # Build a python-docx paragraph wrapper
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, paragraph._parent)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    return p


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--docx",
        type=Path,
        default=Path(r"C:\Users\merao\OneDrive\Documents\Reasearch\Maraol_Feye_Thesis_AASTU_Guideline_Compliant.docx"),
    )
    parser.add_argument("--figures", type=Path, default=ROOT / "out")
    parser.add_argument(
        "--out",
        type=Path,
        default=ROOT / "Maraol_Feye_Thesis_with_Python_Figures.docx",
    )
    args = parser.parse_args()

    if not args.docx.is_file():
        # fallbacks
        candidates = list(Path(r"C:\Users\merao\OneDrive\Documents\Reasearch").glob("Maraol_Feye_Thesis*.docx"))
        candidates = [c for c in candidates if "with_Python" not in c.name and not c.name.endswith(".bak")]
        if not candidates:
            raise SystemExit(f"Thesis DOCX not found: {args.docx}")
        args.docx = sorted(candidates, key=lambda p: p.stat().st_mtime, reverse=True)[0]
        print(f"Using newest thesis: {args.docx.name}")

    doc = Document(str(args.docx))
    caption_re = re.compile(r"^Figure\s+(\d+\.\d+)\b", re.I)
    inserted = []
    seen = set()
    in_lof = False

    for p in list(doc.paragraphs):
        text = (p.text or "").strip()
        # Skip List of Figures section (captions listed as text only)
        style = (p.style.name if p.style is not None else "") or ""
        if style.startswith("Heading") and "list of figures" in text.lower():
            in_lof = True
            continue
        if style.startswith("Heading") and in_lof:
            in_lof = False  # next real chapter heading ends LoF
        if in_lof:
            continue

        m = caption_re.match(text)
        if not m:
            continue
        fig_id = m.group(1)
        if fig_id in seen:
            continue
        fname = FIG_MAP.get(fig_id)
        if not fname:
            continue
        img = args.figures / fname
        if not img.is_file():
            print(f"  skip Figure {fig_id}: missing {img.name}")
            continue
        insert_picture_before(p, img, width_in=5.9)
        seen.add(fig_id)
        inserted.append(fig_id)
        print(f"  inserted Figure {fig_id} ← {fname}")

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"Saved: {args.out}")
    print(f"Inserted {len(inserted)} figures: {inserted}")
    print("Open the DOCX and remove any leftover [FIGURE PLACEHOLDER] boxes.")


if __name__ == "__main__":
    main()
